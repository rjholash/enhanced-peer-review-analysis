#!/usr/bin/env python3

"""
You need to run enhanced _gui_launcher.py to use this script!
Enhanced Instructor peer‑evaluation report generator with reliability metrics
Includes inter-rater reliability (ICC) and Cronbach's alpha calculations
"""

import argparse
import html
import logging
from pathlib import Path
import re
import pandas as pd
import numpy as np
from jinja2 import Template
from scipy import stats
import warnings
warnings.filterwarnings('ignore')

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s ‑ %(levelname)s ‑ %(message)s",
)
LOG = logging.getLogger(__name__)


# ──────────────────────────────── Reliability Calculations ──────────────────────────────── #

def calculate_cronbach_alpha(data: pd.DataFrame) -> float:
    """
    Calculate Cronbach's alpha for internal consistency.
    
    Args:
        data: DataFrame with items as columns and observations as rows
        
    Returns:
        Cronbach's alpha value (0-1, higher is better consistency)
    """
    if data.shape[1] < 2:
        return np.nan
    
    # Remove rows with any missing values
    data_clean = data.dropna()
    if data_clean.empty or data_clean.shape[0] < 2:
        return np.nan
    
    k = data_clean.shape[1]  # number of items
    item_vars = data_clean.var(axis=0, ddof=1)
    total_var = data_clean.sum(axis=1).var(ddof=1)
    
    if total_var == 0:
        return np.nan
    
    alpha = (k / (k - 1)) * (1 - item_vars.sum() / total_var)
    return max(0, alpha)  # Alpha can be negative, but we'll floor at 0


def calculate_icc(data: pd.DataFrame, icc_type: str = 'ICC(2,1)') -> tuple[float, float]:
    """
    Calculate Intraclass Correlation Coefficient for inter-rater reliability.
    
    Args:
        data: DataFrame with raters as columns and subjects as rows
        icc_type: Type of ICC to calculate
        
    Returns:
        Tuple of (ICC value, F-statistic p-value)
    """
    if data.shape[1] < 2 or data.shape[0] < 2:
        return np.nan, np.nan
    
    # Remove rows with any missing values
    data_clean = data.dropna()
    if data_clean.empty or data_clean.shape[0] < 2:
        return np.nan, np.nan
    
    n_subjects = data_clean.shape[0]
    n_raters = data_clean.shape[1]
    
    # Calculate means
    subject_means = data_clean.mean(axis=1)
    rater_means = data_clean.mean(axis=0)
    grand_mean = data_clean.values.mean()
    
    # Sum of squares calculations
    ss_total = np.sum((data_clean.values - grand_mean) ** 2)
    ss_between_subjects = n_raters * np.sum((subject_means - grand_mean) ** 2)
    ss_between_raters = n_subjects * np.sum((rater_means - grand_mean) ** 2)
    ss_error = ss_total - ss_between_subjects - ss_between_raters
    
    # Mean squares
    ms_between_subjects = ss_between_subjects / (n_subjects - 1)
    ms_between_raters = ss_between_raters / (n_raters - 1)
    ms_error = ss_error / ((n_subjects - 1) * (n_raters - 1))
    
    # ICC(2,1) - Two-way random effects, single measurement, absolute agreement
    if ms_error == 0:
        return np.nan, np.nan
    
    icc = (ms_between_subjects - ms_error) / (ms_between_subjects + (n_raters - 1) * ms_error)
    
    # F-test for significance
    f_stat = ms_between_subjects / ms_error
    df1 = n_subjects - 1
    df2 = (n_subjects - 1) * (n_raters - 1)
    p_value = 1 - stats.f.cdf(f_stat, df1, df2)
    
    return max(0, icc), p_value


def calculate_reviewer_scores(merged_df: pd.DataFrame, rating_cols: list) -> pd.DataFrame:
    """
    Calculate reviewer quality scores based on inter-rater reliability metrics.
    Now includes review count tracking and more nuanced bias assessment.
    
    Args:
        merged_df: DataFrame with review data
        rating_cols: List of column names containing ratings
        
    Returns:
        DataFrame with reviewer scores and metrics
    """
    reviewer_scores = []
    
    # Group by the item being reviewed (each group presentation)
    for group_reviewed, group_data in merged_df.groupby('Group#_reviewing'):
        if len(group_data) < 2:  # Need at least 2 reviewers
            continue
            
        # Create matrix: reviewers × rating dimensions
        reviewer_ids = group_data['email_clean'].tolist()
        rating_matrix = group_data[rating_cols].values
        review_times = group_data['DurationMin'].values
        
        # Calculate Cronbach's alpha for this group's reviews
        alpha = calculate_cronbach_alpha(pd.DataFrame(rating_matrix))
        
        # Calculate ICC for inter-rater reliability
        # Transpose so subjects are rows, raters are columns
        if len(rating_cols) >= 2:
            icc_data = pd.DataFrame(rating_matrix.T, columns=reviewer_ids)
            icc_value, icc_p = calculate_icc(icc_data)
        else:
            icc_value, icc_p = np.nan, np.nan
        
        # Calculate individual reviewer deviation from group mean
        group_means = np.mean(rating_matrix, axis=0)
        overall_group_mean = np.mean(group_means)
        
        for i, reviewer_id in enumerate(reviewer_ids):
            reviewer_ratings = rating_matrix[i, :]
            review_time = review_times[i]
            
            # Calculate reviewer-specific metrics
            deviation_from_mean = np.mean(np.abs(reviewer_ratings - group_means))
            consistency_score = 1 / (1 + deviation_from_mean)  # Higher is better
            
            # Apply time penalty only for very short reviews (but don't penalize reasonable variation)
            if review_time < 1.0:  # Per review, not total
                consistency_score *= 0.7  # Reduce consistency score for rushed individual reviews
            
            # Severity bias (tendency to rate higher or lower than group)
            severity_bias = np.mean(reviewer_ratings) - overall_group_mean
            
            # Halo effect (variance in ratings - lower variance suggests halo effect)
            halo_score = np.var(reviewer_ratings) if np.var(reviewer_ratings) > 0 else 0
            
            reviewer_scores.append({
                'reviewer_email': reviewer_id,
                'group_reviewed': group_reviewed,
                'cronbach_alpha': alpha,
                'icc_value': icc_value,
                'icc_p_value': icc_p,
                'consistency_score': consistency_score,
                'severity_bias': severity_bias,
                'halo_score': halo_score,
                'deviation_from_mean': deviation_from_mean,
                'review_time': review_time
            })
    
    scores_df = pd.DataFrame(reviewer_scores)
    
    if not scores_df.empty:
        # Calculate overall reviewer quality scores
        reviewer_summary = scores_df.groupby('reviewer_email').agg({
            'consistency_score': 'mean',
            'severity_bias': lambda x: np.abs(x).mean(),  # Average absolute bias
            'halo_score': 'mean',
            'deviation_from_mean': 'mean',
            'cronbach_alpha': 'mean',
            'icc_value': 'mean',
            'review_time': 'sum'  # Total time across all reviews
        }).reset_index()
        
        # Add review count
        review_counts = scores_df.groupby('reviewer_email').size().reset_index(name='review_count')
        reviewer_summary = reviewer_summary.merge(review_counts, on='reviewer_email')
        
        # Calculate z-scores for severity bias to identify truly extreme outliers
        if len(reviewer_summary) > 1:
            severity_values = scores_df.groupby('reviewer_email')['severity_bias'].mean()
            severity_mean = severity_values.mean()
            severity_std = severity_values.std()
            if severity_std > 0:
                reviewer_summary['severity_bias_raw'] = severity_values.values
                reviewer_summary['severity_bias_zscore'] = (severity_values - severity_mean) / severity_std
            else:
                reviewer_summary['severity_bias_raw'] = severity_values.values
                reviewer_summary['severity_bias_zscore'] = 0
        else:
            reviewer_summary['severity_bias_raw'] = reviewer_summary['severity_bias']
            reviewer_summary['severity_bias_zscore'] = 0
        
        # Composite reliability score (0-100 scale)
        # Weight: 40% consistency, 30% low bias, 20% appropriate variance, 10% ICC
        # Apply penalty only for incomplete reviews or extreme statistical outliers
        incomplete_penalty = reviewer_summary['review_count'].apply(lambda x: 0.5 if x < 2 else 1.0)
        extreme_bias_penalty = reviewer_summary['severity_bias_zscore'].apply(
            lambda x: 0.8 if abs(x) > 2.0 else 1.0  # Only penalize extreme outliers (>2 SD)
        )
        total_time_penalty = reviewer_summary['review_time'].apply(
            lambda x: 0.7 if x < 2.0 else 1.0  # Penalty for insufficient total time for 2 reviews
        )
        
        reviewer_summary['reliability_score'] = (
            reviewer_summary['consistency_score'] * 40 +
            (1 / (1 + reviewer_summary['severity_bias'])) * 30 +
            np.clip(reviewer_summary['halo_score'] * 10, 0, 20) +  # Cap halo contribution
            reviewer_summary['icc_value'].fillna(0.5) * 10
        ) * incomplete_penalty * extreme_bias_penalty * total_time_penalty
        
        return reviewer_summary
    else:
        return pd.DataFrame()

# ──────────────────────────────── Helpers ──────────────────────────────── #

def clean(series: pd.Series, pad: int | None = None) -> pd.Series:
    """Strip, lowercase; optionally zero‑pad numeric IDs."""
    s = series.astype(str).str.strip().str.lower()
    if pad:
        s = s.str.zfill(pad)
    return s


def _load_table(path: Path) -> pd.DataFrame:
    """Load spreadsheet-like input from Excel or CSV."""
    suffix = path.suffix.lower()
    if suffix == ".csv":
        return pd.read_csv(path)
    if suffix in {".xlsx", ".xls"}:
        return pd.read_excel(path)
    raise ValueError(f"Unsupported file type: {path.suffix}. Use .xlsx, .xls, or .csv.")


def _normalize_header(header: object) -> str:
    text = str(header).strip().lower()
    return re.sub(r"[^a-z0-9]+", "", text)


def _find_column(df: pd.DataFrame, label: str, candidates: list[str]) -> str:
    """Match a column by normalized header names."""
    normalized = {_normalize_header(col): str(col) for col in df.columns}
    for candidate in candidates:
        match = normalized.get(_normalize_header(candidate))
        if match:
            return match
    raise KeyError(
        f"Could not find a column for {label}. "
        f"Expected one of: {', '.join(candidates)}. "
        f"Found columns: {', '.join(map(str, df.columns))}"
    )


def read_review(path: Path) -> pd.DataFrame:
    df = _load_table(path)
    email_col = _find_column(df, "review email", ["Email", "E-mail", "Email Address"])
    student_id_col = _find_column(df, "review student ID", ["StudentID", "Student ID", "OrgDefinedId"])
    df["email_clean"] = clean(df[email_col])
    df["student_id_clean"] = clean(df[student_id_col])
    return df


def read_roster(path: Path) -> pd.DataFrame:
    df = _load_table(path)
    email_col = _find_column(df, "roster email", ["Email", "E-mail", "Email Address"])
    student_id_col = _find_column(df, "roster student ID", ["OrgDefinedId", "Org Defined ID", "StudentID", "Student ID"])
    group_col = _find_column(
        df,
        "roster group",
        [
            "Product Review Groups",
            "Product Review Group",
            "Produce Review Groups",
            "Produce Review Group",
            "Group",
            "Groups",
            "Group Name",
        ],
    )

    df["email_clean"] = clean(df[email_col])
    df["student_id_clean"] = clean(df[student_id_col])

    group_text = df[group_col].astype(str).str.strip()
    extracted_groups = group_text.str.extract(r"(\d{1,3})(?!.*\d)", expand=False)
    df["Group"] = extracted_groups.fillna(group_text).str.zfill(2)
    return df


RATING_COLUMNS = [
    "Video_Quality", "Presenters", "Explanation",
    "Mechanism", "Side_Effects", "Bias",
    "Critical_review", "Study_Quality", "Study_participants",
]

PEER_REVIEW_COMPLETION_WEIGHT = 25.0
PEER_SCORE_WEIGHT = 35.0
WRITTEN_REPORT_WEIGHT = 40.0


def _normalize_group_value(value: object) -> str:
    text = str(value).strip()
    matches = re.findall(r"(\d{1,3})", text)
    if matches:
        return matches[-1].zfill(2)
    return text.lower()


def reliability_score_to_grade(score, total_time_minutes, num_reviews, severity_bias_zscore):
    """Convert reliability score (0-100) to grade out of 10."""
    base_grade = 7.0

    if pd.isna(score) or score == 'N/A':
        base_grade = 7.0
    elif score >= 85:
        base_grade = 10.0
    elif score >= 75:
        base_grade = 9.0
    elif score >= 65:
        base_grade = 8.0
    elif score >= 55:
        base_grade = 7.0
    elif score >= 45:
        base_grade = 6.0
    elif score >= 35:
        base_grade = 5.0
    else:
        base_grade = 4.0

    if num_reviews < 2:
        base_grade = base_grade * 0.5
        return base_grade, "incomplete", -50

    time_adjustment = 0.0
    time_category = ""

    if total_time_minutes > 20.0:
        time_adjustment = +1.0
        time_category = "thorough"
    elif total_time_minutes > 15.0:
        time_adjustment = +0.5
        time_category = "good"
    elif total_time_minutes > 10.0:
        time_adjustment = 0.0
        time_category = "adequate"
    elif total_time_minutes > 5.0:
        time_adjustment = -0.5
        time_category = "somewhat_rushed"
    elif total_time_minutes > 2.0:
        time_adjustment = -1.0
        time_category = "rushed"
    elif total_time_minutes > 0.5:
        time_adjustment = -2.0
        time_category = "very_rushed"
    else:
        time_adjustment = -3.0
        time_category = "technical_issue"

    base_grade = min(10.0, max(2.0, base_grade + time_adjustment))

    if not pd.isna(severity_bias_zscore) and abs(severity_bias_zscore) > 2.0:
        base_grade = max(3.0, base_grade - 1.5)

    return base_grade, time_category, time_adjustment


def moderate_peer_review_completion_grade(raw_grade: float, num_reviews: int) -> float:
    """Reduce the impact of low peer-review completion grades without changing published weights."""
    if pd.isna(raw_grade):
        return raw_grade
    if num_reviews < 2:
        return round(raw_grade, 2)

    adjusted_grade = 5.0 + 0.5 * raw_grade
    return round(min(10.0, max(raw_grade, adjusted_grade)), 2)


def _format_special_notes(total_time, num_reviews, severity_bias, severity_bias_zscore) -> list[str]:
    special_notes: list[str] = []
    if num_reviews < 2:
        special_notes.append(f"Incomplete: Only {num_reviews} review(s) completed (required: 2).")
    elif total_time < 2.0:
        special_notes.append(f"Time warning: Total time for 2 reviews was {total_time:.1f} minutes.")

    if not pd.isna(severity_bias_zscore) and abs(severity_bias_zscore) > 2.0:
        special_notes.append(
            f"Extreme bias pattern: {severity_bias_zscore:.1f} standard deviations from peer average."
        )
    elif not pd.isna(severity_bias) and abs(severity_bias) > 0.5:
        bias_type = "higher" if severity_bias > 0 else "lower"
        special_notes.append(f"Rating tendency: rates {bias_type} than peers, but not penalized.")
    return special_notes


def get_feedback_message(score, grade, total_time_minutes, num_reviews, severity_bias, severity_bias_zscore, time_category, time_adjustment):
    """Generate personalized feedback with improved time messaging."""
    if num_reviews < 2:
        if num_reviews == 0:
            return (
                "No peer reviews were submitted, so no peer review completion credit could be awarded. "
                "Because this assignment required two completed reviews, your peer review completion grade for this component is 0. "
                "Please ensure you complete all required reviews in future assignments."
            )
        completion_warning = (
            f" Important: you completed only {num_reviews} review(s) instead of the required 2. "
            "This reduced your peer review completion mark."
        )
        return (
            "Your peer review completion has been recorded."
            f"{completion_warning} Please ensure you complete all required reviews in future assignments."
        )

    time_feedback = ""
    if time_category == "thorough":
        time_feedback = (
            f" Excellent: you spent {total_time_minutes:.1f} minutes on your reviews, "
            "demonstrating exceptional thoroughness. This earned a +1.0 bonus point."
        )
    elif time_category == "good":
        time_feedback = (
            f" Great work: you spent {total_time_minutes:.1f} minutes on your reviews, "
            "showing good thoroughness. This earned a +0.5 bonus point."
        )
    elif time_category == "adequate":
        time_feedback = (
            f" Good: you spent {total_time_minutes:.1f} minutes on your reviews, "
            "which is appropriate for thoughtful evaluation."
        )
    elif time_category == "somewhat_rushed":
        time_feedback = (
            f" Note: you spent {total_time_minutes:.1f} minutes on your reviews. "
            "While acceptable, taking a bit more time could improve the depth of your evaluations (-0.5 points)."
        )
    elif time_category == "rushed":
        time_feedback = (
            f" Feedback: you spent {total_time_minutes:.1f} minutes on your reviews. "
            "For more thorough peer evaluation, consider spending 1-2 minutes per review section (-1.0 point)."
        )
    elif time_category == "very_rushed":
        time_feedback = (
            f" Concern: you spent {total_time_minutes:.1f} minutes on your reviews. "
            "This seems quite rushed for thoughtful evaluation. Please take more time to provide meaningful feedback (-2.0 points)."
        )
    elif time_category == "technical_issue":
        time_feedback = (
            f" Please contact me: your recorded time was {total_time_minutes:.1f} minutes "
            f"({total_time_minutes*60:.0f} seconds). This seems unusually short and may indicate a technical issue. "
            "Please email me to discuss this and I can adjust your grade if there was a technical problem "
            "(-3.0 points pending review)."
        )

    bias_feedback = ""
    if not pd.isna(severity_bias) and abs(severity_bias) > 0.3:
        if severity_bias > 0.3:
            if not pd.isna(severity_bias_zscore) and severity_bias_zscore > 2.0:
                bias_feedback = (
                    f" Rating pattern: your reviews consistently rate higher than peers "
                    f"(z-score: {severity_bias_zscore:.1f}). Consider whether you are applying the criteria as intended. "
                    "This extreme pattern resulted in a -1.5 point adjustment."
                )
            else:
                bias_feedback = (
                    f" Rating observation: your reviews tend to rate higher than peers "
                    f"(average difference: +{severity_bias:.1f}). Consider using the full range of the rating scale. No penalty applied."
                )
        else:
            if not pd.isna(severity_bias_zscore) and severity_bias_zscore < -2.0:
                bias_feedback = (
                    f" Rating pattern: your reviews consistently rate lower than peers "
                    f"(z-score: {severity_bias_zscore:.1f}). This extreme pattern resulted in a -1.5 point adjustment."
                )
            else:
                bias_feedback = (
                    f" Rating observation: your reviews tend to rate lower than peers "
                    f"(average difference: {severity_bias:.1f}). High standards are valuable, but make sure strong work is recognized when present. No penalty applied."
                )

    if pd.isna(score) or score == 'N/A':
        base_message = "Your peer review completion has been recorded. There was insufficient data for detailed reliability analysis."
    elif score >= 85:
        base_message = "Outstanding work. Your reviews demonstrate excellent reliability and strong alignment with peer consensus."
    elif score >= 75:
        base_message = "Excellent work. Your reviews show strong reliability and good consistency with peer consensus."
    elif score >= 65:
        base_message = "Good work. Your reviews show reasonable consistency with peer consensus."
    elif score >= 55:
        base_message = "Satisfactory work. Your reviews show moderate consistency with peers."
    elif score >= 45:
        base_message = "Your reviews show some consistency challenges. Consider using more of the rating scale range and referring back to the evaluation guidelines."
    else:
        base_message = "Your reviews show significant variance from peer consensus. Please review the evaluation guidelines carefully."

    support_message = ""
    if score < 55 or (not pd.isna(severity_bias_zscore) and abs(severity_bias_zscore) > 2.0) or time_category in ["very_rushed", "technical_issue"]:
        support_message = " If you have questions about this assessment or need clarification on the evaluation criteria, please contact me."

    return base_message + time_feedback + bias_feedback + support_message


def build_student_feedback_fields(student_results: pd.DataFrame) -> pd.DataFrame:
    """Attach plain-text feedback fields suitable for mail merge."""
    feedback_messages = []
    note_texts = []
    summary_sentences = []

    for _, row in student_results.iterrows():
        score = row.get("reliability_score", np.nan)
        total_time = row.get("total_time_minutes", 0.0)
        num_reviews = row.get("num_reviews", 0)
        severity_bias = row.get("severity_bias", np.nan)
        severity_bias_zscore = row.get("severity_bias_zscore", np.nan)
        grade = row.get("peer_review_grade_10", 0.0)
        time_category = row.get("time_category", "")
        time_adjustment = row.get("time_adjustment", 0.0)

        feedback = get_feedback_message(
            score, grade, total_time, num_reviews, severity_bias, severity_bias_zscore, time_category, time_adjustment
        )
        notes = _format_special_notes(total_time, num_reviews, severity_bias, severity_bias_zscore)

        summary = (
            (
                "You did not complete any of the 2 required peer reviews, so your peer review completion "
                f"grade was {grade:.1f}/10, contributing {row.get('peer_review_completion_component_25', 0.0):.2f}/25 to the assignment."
                if num_reviews == 0
                else f"Your peer review completion grade was {grade:.1f}/10, contributing "
                f"{row.get('peer_review_completion_component_25', 0.0):.2f}/25 to the assignment. "
                f"You completed {num_reviews}/2 required reviews and spent {total_time:.1f} minutes in total."
            )
        )

        feedback_messages.append(feedback)
        note_texts.append(" | ".join(notes))
        summary_sentences.append(summary)

    student_results["peer_review_feedback"] = feedback_messages
    student_results["peer_review_notes"] = note_texts
    student_results["peer_review_summary"] = summary_sentences
    return student_results


def _build_analysis_frames(review: pd.DataFrame, roster: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    """Build merged review data, group stats, and reviewer summary."""
    merged = review.merge(
        roster[["email_clean", "student_id_clean", "Group"]],
        on="email_clean",
        how="left",
    )
    merged["DurationMin"] = (
        pd.to_datetime(merged["Completion time"]) -
        pd.to_datetime(merged["Start time"])
    ).dt.total_seconds() / 60
    merged["OverallPct"] = merged[RATING_COLUMNS].mean(axis=1) * 10

    reviewer_scores = calculate_reviewer_scores(merged, RATING_COLUMNS)

    group_stats = (
        merged.groupby("Group#_reviewing")["OverallPct"]
        .agg(Mean="mean", SD="std", Reviews="count")
        .round({"Mean": 1, "SD": 1})
        .reset_index(names="Group")
        .sort_values("Group")
    )
    group_stats["GroupKey"] = group_stats["Group"].apply(_normalize_group_value)
    group_stats["Peer Score Component (35%)"] = (group_stats["Mean"] * (PEER_SCORE_WEIGHT / 100.0)).round(2)

    return merged, group_stats, reviewer_scores


def build_individual_peer_review_grades(review: pd.DataFrame, roster: pd.DataFrame) -> pd.DataFrame:
    """Return one row per student with individual peer-review completion grades."""
    merged, _, reviewer_scores = _build_analysis_frames(review, roster)

    reviewer_stats = (
        merged.groupby("email_clean")
        .agg(total_time_minutes=("DurationMin", "sum"), num_reviews=("Group#_reviewing", "count"))
        .reset_index()
    )

    student_results = roster.copy()
    student_results = student_results.merge(reviewer_stats, on="email_clean", how="left")
    student_results = student_results.merge(
        reviewer_scores[[
            "reviewer_email",
            "reliability_score",
            "consistency_score",
            "severity_bias",
            "severity_bias_zscore",
            "cronbach_alpha",
            "icc_value",
        ]],
        left_on="email_clean",
        right_on="reviewer_email",
        how="left",
    )

    student_results["total_time_minutes"] = student_results["total_time_minutes"].fillna(0.0)
    student_results["num_reviews"] = student_results["num_reviews"].fillna(0).astype(int)

    if "severity_bias_zscore" not in student_results.columns:
        student_results["severity_bias_zscore"] = np.nan

    grade_values = []
    time_categories = []
    time_adjustments = []
    for _, row in student_results.iterrows():
        grade, time_category, time_adjustment = reliability_score_to_grade(
            row.get("reliability_score", np.nan),
            row.get("total_time_minutes", 0.0),
            row.get("num_reviews", 0),
            row.get("severity_bias_zscore", np.nan),
        )
        if row.get("num_reviews", 0) == 0:
            grade = 0.0
        grade_values.append(round(grade, 2))
        time_categories.append(time_category)
        time_adjustments.append(time_adjustment)

    student_results["raw_peer_review_grade_10"] = grade_values
    student_results["peer_review_grade_10"] = [
        moderate_peer_review_completion_grade(grade, num_reviews)
        for grade, num_reviews in zip(grade_values, student_results["num_reviews"])
    ]
    student_results["time_category"] = time_categories
    student_results["time_adjustment"] = time_adjustments
    student_results["peer_review_completion_component_25"] = (
        student_results["peer_review_grade_10"] / 10.0 * PEER_REVIEW_COMPLETION_WEIGHT
    ).round(2)

    student_results["GroupKey"] = student_results["Group"].apply(_normalize_group_value)

    name_columns = [col for col in roster.columns if "name" in col.lower()]
    if len(name_columns) >= 2:
        first_name_col = next((col for col in name_columns if "first" in col.lower()), name_columns[0])
        last_name_col = next((col for col in name_columns if "last" in col.lower()), name_columns[1])
        student_results["StudentName"] = (
            student_results[first_name_col].fillna("").astype(str).str.strip() + " " +
            student_results[last_name_col].fillna("").astype(str).str.strip()
        ).str.strip()
    elif len(name_columns) == 1:
        student_results["StudentName"] = student_results[name_columns[0]].fillna("").astype(str).str.strip()
    else:
        student_results["StudentName"] = ""

    return student_results


def _extract_written_grade_token(text: str) -> str | None:
    labels = [
        "Written Report Grade",
        "Report Grade",
        "Written Grade",
        "Final Grade",
        "Grade",
    ]
    token_pattern = r"(\d+(?:\.\d+)?(?:\s*/\s*\d+(?:\.\d+)?)?\s*%?)"
    for label in labels:
        match = re.search(rf"{re.escape(label)}[^0-9]{{0,50}}{token_pattern}", text, flags=re.IGNORECASE)
        if match:
            return match.group(1).strip()
    fallback = re.search(r"(\d+(?:\.\d+)?)\s*/\s*40", text, flags=re.IGNORECASE)
    if fallback:
        return f"{fallback.group(1)}/40"
    return None


def _parse_grade_token(token: object) -> tuple[float, float]:
    text = str(token).strip()
    if not text or text.lower() == "nan":
        return np.nan, np.nan

    slash_match = re.search(r"(\d+(?:\.\d+)?)\s*/\s*(\d+(?:\.\d+)?)", text)
    if slash_match:
        score = float(slash_match.group(1))
        maximum = float(slash_match.group(2))
        if maximum == 0:
            return np.nan, np.nan
        percent = score / maximum * 100.0
        component = score if maximum == 40 else percent * (WRITTEN_REPORT_WEIGHT / 100.0)
        return round(percent, 2), round(component, 2)

    percent_match = re.search(r"(\d+(?:\.\d+)?)\s*%", text)
    if percent_match:
        percent = float(percent_match.group(1))
        return round(percent, 2), round(percent * (WRITTEN_REPORT_WEIGHT / 100.0), 2)

    numeric_match = re.search(r"\d+(?:\.\d+)?", text)
    if not numeric_match:
        return np.nan, np.nan

    value = float(numeric_match.group(0))
    if value <= 1.0:
        percent = value * 100.0
        return round(percent, 2), round(percent * (WRITTEN_REPORT_WEIGHT / 100.0), 2)
    if value <= 40.0:
        return round(value / 40.0 * 100.0, 2), round(value, 2)
    return round(value, 2), round(value * (WRITTEN_REPORT_WEIGHT / 100.0), 2)


def load_written_report_grades(source: Path) -> pd.DataFrame:
    """Load group written-report grades from a table file or a folder of feedback_group HTML files."""
    rows = []

    if source.is_dir():
        html_files = sorted(source.glob("feedback_group*.html"))
        if not html_files:
            raise ValueError(f"No feedback_group*.html files found in {source}")
        for html_file in html_files:
            group_match = re.search(r"group[-_]?(\d+)", html_file.stem, flags=re.IGNORECASE)
            if not group_match:
                continue
            token = _extract_written_grade_token(html_file.read_text(encoding="utf-8", errors="ignore"))
            if not token:
                raise ValueError(f"Could not find a written report grade inside {html_file.name}")
            percent, component = _parse_grade_token(token)
            rows.append({
                "GroupKey": group_match.group(1).zfill(2),
                "written_report_source": html_file.name,
                "written_report_raw": token,
                "written_report_percent": percent,
                "written_report_component_40": component,
            })
    else:
        df = _load_table(source)
        group_col = _find_column(
            df,
            "written report group",
            ["Group", "Product Review Groups", "Produce Review Groups", "Group Name"],
        )
        grade_col = _find_column(
            df,
            "written report grade",
            ["Written Report Grade", "Report Grade", "Written Grade", "Grade", "Score", "Percent", "Percentage"],
        )
        for _, row in df.iterrows():
            percent, component = _parse_grade_token(row[grade_col])
            rows.append({
                "GroupKey": _normalize_group_value(row[group_col]),
                "written_report_source": source.name,
                "written_report_raw": row[grade_col],
                "written_report_percent": percent,
                "written_report_component_40": component,
            })

    result = pd.DataFrame(rows)
    if result.empty:
        raise ValueError("No written report grades could be loaded.")

    result = result.drop_duplicates(subset=["GroupKey"], keep="first")
    return result


def load_enhanced_group_report_summaries(source: Path) -> pd.DataFrame:
    """Load summary metrics from group peer review report HTML files."""
    if not source.is_dir():
        raise ValueError("Enhanced group report source must be a folder.")

    rows = []
    html_files = sorted(source.glob("enhanced_group_*_report.html")) + sorted(source.glob("Group_*Peer_Review_Report.html")) + sorted(source.glob("Group_*_Peer_Review_Report.html"))
    if not html_files:
        raise ValueError(f"No group peer review report HTML files found in {source}")

    for html_file in html_files:
        group_match = re.search(r"(?:enhanced_group_(\d+)_report|group_(\d+)peer_review_report|group_(\d+)_peer_review_report)", html_file.stem, flags=re.IGNORECASE)
        if not group_match:
            continue
        text = html_file.read_text(encoding="utf-8", errors="ignore")
        peer_grade_match = re.search(r"Peer Review Grade:\s*([\d.]+)%", text, flags=re.IGNORECASE)
        component_35_match = re.search(
            r"(?:Peer Review Component|Weighted Score) \((?:35|50|40)%\):\s*([\d.]+)(?:/35|/50|%)",
            text,
            flags=re.IGNORECASE,
        )
        total_mark_match = re.search(r"Total D2L Mark:\s*([\d.]+)/60", text, flags=re.IGNORECASE)
        group_number = next(value for value in group_match.groups() if value is not None)
        rows.append({
            "GroupKey": group_number.zfill(2),
            "enhanced_group_report_source": html_file.name,
            "group_peer_review_percent_from_report": float(peer_grade_match.group(1)) if peer_grade_match else np.nan,
            "group_peer_review_35_from_report": float(component_35_match.group(1)) if component_35_match else np.nan,
            "legacy_group_total_mark_60": float(total_mark_match.group(1)) if total_mark_match else np.nan,
        })

    result = pd.DataFrame(rows)
    if result.empty:
        raise ValueError("No enhanced group report summaries could be loaded.")
    return result.drop_duplicates(subset=["GroupKey"], keep="first")


def build_consolidated_gradebook(
    review: pd.DataFrame,
    roster: pd.DataFrame,
    written_report_source: Path | None = None,
    enhanced_group_report_source: Path | None = None,
) -> pd.DataFrame:
    """Build the detailed per-student gradebook for mail merge and D2L import."""
    student_results = build_individual_peer_review_grades(review, roster)
    student_results = build_student_feedback_fields(student_results)
    _, group_stats, _ = _build_analysis_frames(review, roster)

    consolidated = student_results.merge(
        group_stats[["GroupKey", "Mean", "Peer Score Component (35%)"]],
        on="GroupKey",
        how="left",
    )
    consolidated = consolidated.rename(columns={
        "Mean": "group_peer_score_percent",
        "Peer Score Component (35%)": "peer_scores_component_35",
    })

    if written_report_source is not None:
        written_report = load_written_report_grades(written_report_source)
        consolidated = consolidated.merge(written_report, on="GroupKey", how="left")
    else:
        consolidated["written_report_source"] = ""
        consolidated["written_report_raw"] = ""
        consolidated["written_report_percent"] = np.nan
        consolidated["written_report_component_40"] = np.nan

    if enhanced_group_report_source is not None:
        group_report_summary = load_enhanced_group_report_summaries(enhanced_group_report_source)
        consolidated = consolidated.merge(group_report_summary, on="GroupKey", how="left")
    else:
        consolidated["enhanced_group_report_source"] = ""
        consolidated["group_peer_review_percent_from_report"] = np.nan
        consolidated["group_peer_review_35_from_report"] = np.nan
        consolidated["legacy_group_total_mark_60"] = np.nan

    if "group_peer_review_35_from_report" in consolidated.columns:
        consolidated["peer_scores_component_35"] = consolidated["group_peer_review_35_from_report"].fillna(
            consolidated["peer_scores_component_35"]
        )
    if "group_peer_review_percent_from_report" in consolidated.columns:
        consolidated["group_peer_score_percent"] = consolidated["group_peer_review_percent_from_report"].fillna(
            consolidated["group_peer_score_percent"]
        )

    consolidated["peer_review_completion_component_25"] = consolidated["peer_review_completion_component_25"].fillna(0.0)
    consolidated["peer_scores_component_35"] = consolidated["peer_scores_component_35"].fillna(0.0)
    consolidated["written_report_component_40"] = consolidated["written_report_component_40"].fillna(0.0)
    consolidated["final_assignment_grade_100"] = (
        consolidated["peer_review_completion_component_25"] +
        consolidated["peer_scores_component_35"] +
        consolidated["written_report_component_40"]
    ).round(2)

    export_columns = [
        "student_id_clean",
        "StudentName",
        "email_clean",
        "Group",
        "num_reviews",
        "total_time_minutes",
        "reliability_score",
        "raw_peer_review_grade_10",
        "peer_review_grade_10",
        "peer_review_completion_component_25",
        "peer_review_summary",
        "peer_review_feedback",
        "peer_review_notes",
        "group_peer_score_percent",
        "peer_scores_component_35",
        "written_report_raw",
        "written_report_percent",
        "written_report_component_40",
        "group_peer_review_percent_from_report",
        "group_peer_review_35_from_report",
        "legacy_group_total_mark_60",
        "final_assignment_grade_100",
    ]
    consolidated = consolidated[export_columns].rename(columns={
        "student_id_clean": "OrgDefinedId",
        "email_clean": "Email",
        "Group": "Group",
        "num_reviews": "PeerReviewsCompleted",
        "total_time_minutes": "PeerReviewTimeMinutes",
        "reliability_score": "ReliabilityScore",
        "peer_review_grade_10": "PeerReviewGradeOutOf10",
        "raw_peer_review_grade_10": "RawPeerReviewGradeOutOf10",
        "peer_review_completion_component_25": "PeerReviewCompletion25",
        "peer_review_summary": "PeerReviewSummary",
        "peer_review_feedback": "PeerReviewFeedback",
        "peer_review_notes": "PeerReviewNotes",
        "group_peer_score_percent": "PeerScoresPercent",
        "peer_scores_component_35": "PeerScores35",
        "written_report_raw": "WrittenReportRaw",
        "written_report_percent": "WrittenReportPercent",
        "written_report_component_40": "WrittenReport40",
        "group_peer_review_percent_from_report": "GroupPeerReviewPercentFromReport",
        "group_peer_review_35_from_report": "GroupPeerReview35FromReport",
        "legacy_group_total_mark_60": "LegacyGroupTotalMark60",
        "final_assignment_grade_100": "FinalGrade100",
    })
    consolidated = consolidated.sort_values(["Group", "OrgDefinedId", "Email"]).reset_index(drop=True)
    return consolidated


def build_d2l_grade_import(consolidated: pd.DataFrame, grade_item_name: str = "Final Assignment Grade") -> pd.DataFrame:
    """Build a minimal D2L-friendly grade import CSV."""
    return pd.DataFrame({
        "Org Defined ID": consolidated["OrgDefinedId"],
        f"{grade_item_name} Points Grade": consolidated["FinalGrade100"],
        "End-of-Line Indicator": "#",
    })


def _safe_text(value: object, fallback: str = "") -> str:
    if pd.isna(value):
        return fallback
    text = str(value).strip()
    return text if text and text.lower() != "nan" else fallback


def build_word_mail_merge_export(consolidated: pd.DataFrame) -> pd.DataFrame:
    """Build a plain-text, Word-friendly mail merge workbook."""
    rows = []
    for _, row in consolidated.iterrows():
        student_name = _safe_text(row.get("StudentName"), "Student")
        email = _safe_text(row.get("Email"))
        student_id = _safe_text(row.get("OrgDefinedId"))
        group = _safe_text(row.get("Group"))
        peer_summary = _safe_text(row.get("PeerReviewSummary"))
        peer_feedback = _safe_text(row.get("PeerReviewFeedback"))
        peer_notes = _safe_text(row.get("PeerReviewNotes"))
        raw_written = _safe_text(row.get("WrittenReportRaw"))

        peer_completion = float(row.get("PeerReviewCompletion25", 0.0) or 0.0)
        peer_scores = float(row.get("PeerScores35", 0.0) or 0.0)
        written_report = float(row.get("WrittenReport40", 0.0) or 0.0)
        final_grade = float(row.get("FinalGrade100", 0.0) or 0.0)
        reviews_completed = int(row.get("PeerReviewsCompleted", 0) or 0)

        paragraph_lines = [
            f"Your final grade for this assignment was {final_grade:.2f}/100.",
            (
                "This was calculated as follows: "
                f"peer review completion {peer_completion:.2f}/25, "
                f"peer scores and ranking {peer_scores:.2f}/35, "
                f"and written report {written_report:.2f}/40."
            ),
        ]
        if peer_summary:
            paragraph_lines.append(f"Peer review completion summary: {peer_summary}")
        if peer_feedback:
            paragraph_lines.append(f"Peer review feedback: {peer_feedback}")
        if peer_notes:
            paragraph_lines.append(f"Notes: {peer_notes}")

        rows.append({
            "StudentName": student_name,
            "Email": email,
            "OrgDefinedId": student_id,
            "Group": group,
            "Greeting": f"Hello {student_name},",
            "FinalGrade100": round(final_grade, 2),
            "FinalGrade100Text": f"{final_grade:.2f}",
            "PeerReviewCompletion25": round(peer_completion, 2),
            "PeerReviewCompletion25Text": f"{peer_completion:.2f}",
            "PeerScores35": round(peer_scores, 2),
            "PeerScores35Text": f"{peer_scores:.2f}",
            "WrittenReport40": round(written_report, 2),
            "WrittenReport40Text": f"{written_report:.2f}",
            "PeerReviewsCompleted": reviews_completed,
            "WrittenReportRaw": raw_written,
            "PeerReviewSummary": peer_summary,
            "PeerReviewFeedback": peer_feedback,
            "PeerReviewNotes": peer_notes,
            "MergeParagraph": " ".join(paragraph_lines),
        })

    return pd.DataFrame(rows)


def _append_merge_field(paragraph, field_name: str) -> None:
    """Append a Word MERGEFIELD to a paragraph."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    begin = OxmlElement("w:r")
    begin_fld = OxmlElement("w:fldChar")
    begin_fld.set(qn("w:fldCharType"), "begin")
    begin.append(begin_fld)
    paragraph._p.append(begin)

    instr = OxmlElement("w:r")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" MERGEFIELD  {field_name}  \\* MERGEFORMAT "
    instr.append(instr_text)
    paragraph._p.append(instr)

    separate = OxmlElement("w:r")
    separate_fld = OxmlElement("w:fldChar")
    separate_fld.set(qn("w:fldCharType"), "separate")
    separate.append(separate_fld)
    paragraph._p.append(separate)

    text_run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = f"«{field_name}»"
    text_run.append(text)
    paragraph._p.append(text_run)

    end = OxmlElement("w:r")
    end_fld = OxmlElement("w:fldChar")
    end_fld.set(qn("w:fldCharType"), "end")
    end.append(end_fld)
    paragraph._p.append(end)


def write_word_mail_merge_template(template_path: Path) -> Path:
    """Write a plain-text Word mail merge template next to the merge workbook."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required to generate the Word mail merge template. "
            "Install it with: pip install python-docx"
        ) from exc

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Peer Review Assignment Results")
    title_run.bold = True
    title_run.font.size = Pt(16)

    intro = doc.add_paragraph()
    intro.add_run(
        "This email was autogenerated to give you personal feedback on your peer review reports, "
        "and to explain your personal mark for this project."
    )

    greeting = doc.add_paragraph()
    _append_merge_field(greeting, "Greeting")

    body = doc.add_paragraph()
    body.add_run("Your final grade for this assignment was ")
    _append_merge_field(body, "FinalGrade100Text")
    body.add_run("/100.")

    summary = doc.add_paragraph()
    summary.add_run("This was calculated as follows: peer review completion ")
    _append_merge_field(summary, "PeerReviewCompletion25Text")
    summary.add_run("/25, peer scores and ranking ")
    _append_merge_field(summary, "PeerScores35Text")
    summary.add_run("/35, and written report ")
    _append_merge_field(summary, "WrittenReport40Text")
    summary.add_run("/40.")

    completion = doc.add_paragraph()
    completion.add_run("Peer review completion summary: ").bold = True
    _append_merge_field(completion, "PeerReviewSummary")

    feedback = doc.add_paragraph()
    feedback.add_run("Peer review feedback: ").bold = True
    _append_merge_field(feedback, "PeerReviewFeedback")

    notes = doc.add_paragraph()
    notes.add_run("Notes: ").bold = True
    _append_merge_field(notes, "PeerReviewNotes")

    d2l_note = doc.add_paragraph()
    d2l_note.add_run(
        "To understand your group marks for your paper and video please check the group dropbox on D2L shortly for detailed reports"
    ).bold = True

    combined = doc.add_paragraph()
    combined.add_run("Single-field option: ").bold = True
    _append_merge_field(combined, "MergeParagraph")

    closing = doc.add_paragraph()
    closing.add_run("If you have questions about any part of this breakdown, please get in touch.")

    doc.save(template_path)
    return template_path


def render_student_email_html(row: pd.Series) -> str:
    """Render an Outlook-safe HTML email body for one student."""
    student_name = html.escape(_safe_text(row.get("StudentName"), "Student"))
    peer_summary = html.escape(_safe_text(row.get("PeerReviewSummary"), ""))
    peer_feedback = html.escape(_safe_text(row.get("PeerReviewFeedback"), ""))
    peer_notes = html.escape(_safe_text(row.get("PeerReviewNotes"), ""))

    marks_rows = [
        ("Peer review completion", f"{row.get('PeerReviewCompletion25', 0.0):.2f}/25"),
        ("Peer scores and ranking", f"{row.get('PeerScores35', 0.0):.2f}/35"),
        ("Written report", f"{row.get('WrittenReport40', 0.0):.2f}/40"),
        ("Final grade", f"{row.get('FinalGrade100', 0.0):.2f}/100"),
    ]
    marks_html = "".join(
        f"""
        <tr>
            <td style="padding:10px 12px;border-bottom:1px solid #e5e7eb;">{html.escape(label)}</td>
            <td style="padding:10px 12px;border-bottom:1px solid #e5e7eb;font-weight:600;text-align:right;">{html.escape(value)}</td>
        </tr>
        """
        for label, value in marks_rows
    )

    raw_written = _safe_text(row.get("WrittenReportRaw"))
    raw_group = _safe_text(row.get("Group"))
    supplemental_lines = []
    if raw_group:
        supplemental_lines.append(f"Group: {raw_group}")
    if raw_written:
        supplemental_lines.append(f"Written report source grade: {raw_written}")
    if row.get("PeerReviewsCompleted", 0):
        supplemental_lines.append(
            f"Peer reviews completed: {int(row.get('PeerReviewsCompleted', 0))}/2"
        )
    supplemental_html = ""
    if supplemental_lines:
        rows_html = "".join(
            f"""
            <tr>
              <td style="padding:4px 0;font-size:14px;line-height:20px;color:#334155;">
                {html.escape(line)}
              </td>
            </tr>
            """
            for line in supplemental_lines
        )
        supplemental_html = f"""
        <table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="border-collapse:collapse;">
          <tr>
            <td style="padding-top:16px;">
              <table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="border-collapse:collapse;background-color:#f8fafc;border:1px solid #dbe5ef;">
                <tr>
                  <td style="padding:14px 16px;">
                    <table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="border-collapse:collapse;">
                      {rows_html}
                    </table>
                  </td>
                </tr>
              </table>
            </td>
          </tr>
        </table>
        """

    notes_html = ""
    if peer_notes:
        notes_html = f"""
        <table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="border-collapse:collapse;">
          <tr>
            <td style="padding-top:16px;">
              <table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="border-collapse:collapse;background-color:#fff7ed;border:1px solid #fed7aa;">
                <tr>
                  <td style="padding:14px 16px;font-size:14px;line-height:21px;color:#9a3412;">
                    <strong>Notes</strong><br>
                    {peer_notes}
                  </td>
                </tr>
              </table>
            </td>
          </tr>
        </table>
        """

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>Peer Review Assignment Results</title>
</head>
<body style="margin:0;padding:0;background-color:#f3f4f6;">
  <table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="border-collapse:collapse;background-color:#f3f4f6;">
    <tr>
      <td align="center" style="padding:24px 12px;">
        <table role="presentation" width="720" cellpadding="0" cellspacing="0" border="0" style="width:720px;max-width:720px;border-collapse:collapse;background-color:#ffffff;border:1px solid #d1d5db;">
          <tr>
            <td style="padding:24px 32px;background-color:#0f4c81;color:#ffffff;font-family:Segoe UI, Arial, sans-serif;">
              <div style="font-size:12px;line-height:16px;letter-spacing:1px;text-transform:uppercase;">Peer Review Analysis</div>
              <div style="padding-top:10px;font-size:28px;line-height:34px;font-weight:700;">Assignment Results</div>
            </td>
          </tr>
          <tr>
            <td style="padding:28px 32px;font-family:Segoe UI, Arial, sans-serif;color:#111827;">
              <p style="margin:0 0 16px 0;font-size:16px;line-height:24px;">Hello {student_name},</p>
              <p style="margin:0 0 18px 0;font-size:16px;line-height:24px;">
                Here is a summary of your marks for this assignment. The table below shows how your final grade was built from the published grading components.
              </p>

              <table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="border-collapse:collapse;border:1px solid #e5e7eb;">
                <tr style="background-color:#f8fafc;">
                  <th align="left" style="padding:12px;border-bottom:1px solid #e5e7eb;font-size:14px;line-height:20px;font-family:Segoe UI, Arial, sans-serif;color:#111827;">Component</th>
                  <th align="right" style="padding:12px;border-bottom:1px solid #e5e7eb;font-size:14px;line-height:20px;font-family:Segoe UI, Arial, sans-serif;color:#111827;">Mark</th>
                </tr>
                {marks_html}
              </table>

              <table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="border-collapse:collapse;">
                <tr>
                  <td style="padding-top:18px;">
                    <table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="border-collapse:collapse;background-color:#eff6ff;border:1px solid #bfdbfe;">
                      <tr>
                        <td style="padding:16px 18px;font-size:14px;line-height:22px;font-family:Segoe UI, Arial, sans-serif;color:#1f2937;">
                          <strong style="color:#1e3a8a;">Peer review completion summary</strong><br>
                          <span>{peer_summary}</span>
                        </td>
                      </tr>
                    </table>
                  </td>
                </tr>
              </table>

              <table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="border-collapse:collapse;">
                <tr>
                  <td style="padding-top:18px;">
                    <table role="presentation" width="100%" cellpadding="0" cellspacing="0" border="0" style="border-collapse:collapse;background-color:#f8fafc;border:1px solid #dbe5ef;">
                      <tr>
                        <td style="padding:16px 18px;font-size:14px;line-height:22px;font-family:Segoe UI, Arial, sans-serif;color:#1f2937;">
                          <strong>Peer review feedback</strong><br>
                          <span>{peer_feedback}</span>
                        </td>
                      </tr>
                    </table>
                  </td>
                </tr>
              </table>

              {notes_html}
              {supplemental_html}

              <p style="margin:24px 0 0 0;font-size:14px;line-height:22px;font-family:Segoe UI, Arial, sans-serif;color:#4b5563;">
                If you have questions about any part of this breakdown, please reply to this message.
              </p>
            </td>
          </tr>
        </table>
      </td>
    </tr>
  </table>
</body>
</html>"""


def build_html_email_merge(consolidated: pd.DataFrame, subject_template: str = "Peer Review Assignment Results") -> pd.DataFrame:
    """Build an email merge export with HTML bodies."""
    email_rows = []
    for _, row in consolidated.iterrows():
        student_name = _safe_text(row.get("StudentName"), "Student")
        subject = subject_template
        html_body = render_student_email_html(row)
        email_rows.append({
            "To": _safe_text(row.get("Email")),
            "Subject": subject,
            "StudentName": student_name,
            "OrgDefinedId": _safe_text(row.get("OrgDefinedId")),
            "FinalGrade100": row.get("FinalGrade100", 0.0),
            "HTMLBody": html_body,
        })
    return pd.DataFrame(email_rows)


def write_email_previews(email_merge: pd.DataFrame, output_dir: Path, limit: int | None = None) -> list[Path]:
    """Write standalone HTML preview files for email bodies."""
    output_dir.mkdir(parents=True, exist_ok=True)
    paths: list[Path] = []
    rows = email_merge.head(limit) if limit is not None else email_merge
    for _, row in rows.iterrows():
        student_id = _safe_text(row.get("OrgDefinedId"), "unknown")
        safe_id = re.sub(r"[^A-Za-z0-9_-]+", "_", student_id)
        preview_path = output_dir / f"email_preview_{safe_id}.html"
        preview_path.write_text(str(row["HTMLBody"]), encoding="utf-8")
        paths.append(preview_path)
    return paths


def build_enhanced_report(review: pd.DataFrame, roster: pd.DataFrame) -> str:
    """Return enhanced HTML report with reliability metrics."""
    # ── merge on email ───────────────────────────────────────────────
    merged = review.merge(
        roster[["email_clean", "student_id_clean", "Group"]],
        on="email_clean",
        how="left",
        indicator=True,
    )

    # ── submission pattern flags ────────────────────────────────────
    counts = merged.groupby("email_clean")["Group#_reviewing"].count()
    one_only = counts[counts == 1].index.tolist()
    three_plus = counts[counts > 2].index.tolist()
    submitted = set(merged["email_clean"])
    no_submit = sorted(set(roster["email_clean"]) - submitted)

    # map email → ID for nicer lists
    email_to_id = roster.set_index("email_clean")["student_id_clean"].to_dict()
    tag = lambda lst: [f"{email_to_id.get(e,'?')} ({e})" for e in lst]

    # ── duration (minutes) ──────────────────────────────────────────
    merged["DurationMin"] = (
        pd.to_datetime(merged["Completion time"]) -
        pd.to_datetime(merged["Start time"])
    ).dt.total_seconds() / 60

    # ── rating columns: numeric only, exclude helper cols ───────────
    rating_cols = [
        "Video_Quality", "Presenters", "Explanation",
        "Mechanism", "Side_Effects", "Bias",
        "Critical_review", "Study_Quality", "Study_participants",
    ]
    merged["OverallPct"] = merged[rating_cols].mean(axis=1) * 10  # 0‑100 %

    # ── ENHANCED: Calculate reliability metrics ─────────────────────
    LOG.info("Calculating inter-rater reliability metrics...")
    reviewer_scores = calculate_reviewer_scores(merged, rating_cols)
    
    # ── group stats: mean % and SD (%), plus count ──────────────────
    group_stats = (
        merged.groupby("Group#_reviewing")["OverallPct"]
        .agg(Mean="mean", SD="std", Reviews="count")
        .round({"Mean": 1, "SD": 1})
        .reset_index(names="Group")
        .sort_values("Group")
    )
    
    # Group peer review contributes 35% of the assignment grade.
    group_stats["Peer Eval (35%)"] = (group_stats["Mean"] * 0.35).round(1)

    # Add reliability metrics to group stats if available
    if not reviewer_scores.empty:
        group_reliability = []
        for group_id in group_stats['Group']:
            group_reviews = merged[merged['Group#_reviewing'] == group_id]
            if len(group_reviews) >= 2:
                group_matrix = group_reviews[rating_cols].values
                alpha = calculate_cronbach_alpha(pd.DataFrame(group_matrix))
                
                # ICC calculation for group
                if len(rating_cols) >= 2:
                    icc_data = pd.DataFrame(group_matrix.T)
                    icc_val, _ = calculate_icc(icc_data)
                else:
                    icc_val = np.nan
                    
                group_reliability.append({
                    'Group': group_id,
                    'Cronbach_Alpha': round(alpha, 3) if not np.isnan(alpha) else 'N/A',
                    'ICC': round(icc_val, 3) if not np.isnan(icc_val) else 'N/A'
                })
            else:
                group_reliability.append({
                    'Group': group_id,
                    'Cronbach_Alpha': 'N/A',
                    'ICC': 'N/A'
                })
        
        reliability_df = pd.DataFrame(group_reliability)
        group_stats = group_stats.merge(reliability_df, on='Group', how='left')

    html_group = group_stats.to_html(
        index=False, classes="table", border=0,
        justify="left", float_format="%.1f"
    )

    # ── Enhanced time table with reliability scores ─────────────────
    time_data = merged[["email_clean", "Group#_reviewing", "DurationMin"]].copy()
    
    # Add reviewer reliability scores if available
    if not reviewer_scores.empty:
        time_data = time_data.merge(
            reviewer_scores[['reviewer_email', 'reliability_score']],
            left_on='email_clean',
            right_on='reviewer_email',
            how='left'
        )
        time_data['reliability_score'] = time_data['reliability_score'].round(1)
    else:
        time_data['reliability_score'] = 'N/A'

    html_time = (
        time_data[["email_clean", "Group#_reviewing", "DurationMin", "reliability_score"]]
        .sort_values("email_clean")
        .rename(columns={
            "email_clean": "Email",
            "Group#_reviewing": "Group",
            "DurationMin": "Minutes",
            "reliability_score": "Reliability Score"
        })
        .to_html(index=False, classes="table", border=0,
                 justify="left", float_format="%.1f")
    )

    # ── Generate individual student feedback sections ───────────────
    def reliability_score_to_grade(score, total_time_minutes, num_reviews, severity_bias_zscore):
        """Convert reliability score (0-100) to grade out of 10, with improved graduated time adjustments."""
        base_grade = 7.0  # Default grade
        
        if pd.isna(score) or score == 'N/A':
            base_grade = 7.0
        elif score >= 85:
            base_grade = 10.0
        elif score >= 75:
            base_grade = 9.0
        elif score >= 65:
            base_grade = 8.0
        elif score >= 55:
            base_grade = 7.0
        elif score >= 45:
            base_grade = 6.0
        elif score >= 35:
            base_grade = 5.0
        else:
            base_grade = 4.0
        
        # Apply penalty for incomplete reviews (less than 2 required)
        if num_reviews < 2:
            base_grade = base_grade * 0.5  # Half marks for incomplete reviews
            return base_grade, "incomplete", -50  # Return early for incomplete reviews
        
        # Improved graduated time-based adjustments with more supportive approach
        time_adjustment = 0.0
        time_category = ""
        
        if total_time_minutes > 20.0:  # Very thorough reviews
            time_adjustment = +1.0
            time_category = "thorough"
        elif total_time_minutes > 15.0:  # Good thoroughness
            time_adjustment = +0.5
            time_category = "good"
        elif total_time_minutes > 10.0:  # Adequate time
            time_adjustment = 0.0
            time_category = "adequate"
        elif total_time_minutes > 5.0:  # Somewhat rushed but acceptable
            time_adjustment = -0.5
            time_category = "somewhat_rushed"
        elif total_time_minutes > 2.0:  # Rushed but still reasonable
            time_adjustment = -1.0
            time_category = "rushed"
        elif total_time_minutes > 0.5:  # Very rushed - significant concern
            time_adjustment = -2.0
            time_category = "very_rushed"
        else:  # Extremely short - likely technical issue
            time_adjustment = -3.0
            time_category = "technical_issue"
        
        # Apply time adjustment
        base_grade = min(10.0, max(2.0, base_grade + time_adjustment))
        
        # Apply penalty only for extreme bias (>2 standard deviations)
        if not pd.isna(severity_bias_zscore) and abs(severity_bias_zscore) > 2.0:
            base_grade = max(3.0, base_grade - 1.5)  # Reduced from -2.0 to -1.5
        
        return base_grade, time_category, time_adjustment

    def get_feedback_message(score, grade, total_time_minutes, num_reviews, severity_bias, severity_bias_zscore, time_category, time_adjustment):
        """Generate personalized feedback with improved time messaging."""
        
        # Review completion issues
        if num_reviews < 2:
            if num_reviews == 0:
                return (
                    "No peer reviews were submitted, so no peer review completion credit could be awarded. "
                    "Because this assignment required two completed reviews, your peer review completion grade for this component is 0. "
                    "Please ensure you complete all required reviews in future assignments."
                )
            completion_warning = (
                f" <b>Important:</b> You completed only {num_reviews} review(s) instead of the required 2. "
                "This reduced your peer review completion mark."
            )
            return f"Your peer review completion has been recorded.{completion_warning} Please ensure you complete all required reviews in future assignments."
        
        # Time-based feedback with more graduated and supportive messaging
        time_feedback = ""
        if time_category == "thorough":
            time_feedback = f" <b>Excellent:</b> You spent {total_time_minutes:.1f} minutes on your reviews, demonstrating exceptional thoroughness. This has earned you a +1.0 bonus point."
        elif time_category == "good":
            time_feedback = f" <b>Great work:</b> You spent {total_time_minutes:.1f} minutes on your reviews, showing good thoroughness. This has earned you a +0.5 bonus point."
        elif time_category == "adequate":
            time_feedback = f" <b>Good:</b> You spent {total_time_minutes:.1f} minutes on your reviews, which is appropriate for thoughtful evaluation."
        elif time_category == "somewhat_rushed":
            time_feedback = f" <b>Note:</b> You spent {total_time_minutes:.1f} minutes on your reviews. While acceptable, taking a bit more time could improve the depth of your evaluations (-0.5 points)."
        elif time_category == "rushed":
            time_feedback = f" <b>Feedback:</b> You spent {total_time_minutes:.1f} minutes on your reviews. For more thorough peer evaluation, consider spending 1-2 minutes per review section (-1.0 point)."
        elif time_category == "very_rushed":
            time_feedback = f" <b>Concern:</b> You spent {total_time_minutes:.1f} minutes on your reviews. This seems quite rushed for thoughtful evaluation. Please take more time to provide meaningful feedback (-2.0 points)."
        elif time_category == "technical_issue":
            time_feedback = f" <b>Please Contact Me:</b> Your recorded time was {total_time_minutes:.1f} minutes ({total_time_minutes*60:.0f} seconds). This seems unusually short and may indicate a technical issue. Please email me to discuss - if there was a technical problem, I will adjust your grade accordingly (-3.0 points pending review)."
        
        # Bias explanation with more educational focus
        bias_feedback = ""
        if not pd.isna(severity_bias) and abs(severity_bias) > 0.3:  # Lower threshold for educational feedback
            if severity_bias > 0.3:
                if not pd.isna(severity_bias_zscore) and severity_bias_zscore > 2.0:
                    bias_feedback = f" <b>Rating Pattern:</b> Your reviews consistently rate higher than peers (z-score: {severity_bias_zscore:.1f}). Consider whether you're applying the evaluation criteria as intended. This extreme pattern has resulted in a -1.5 point adjustment."
                else:
                    bias_feedback = f" <b>Rating Observation:</b> Your reviews tend to rate higher than peers (avg difference: +{severity_bias:.1f}). This may indicate generous scoring - consider the full range of the rating scale. No penalty applied."
            else:
                if not pd.isna(severity_bias_zscore) and severity_bias_zscore < -2.0:
                    bias_feedback = f" <b>Rating Pattern:</b> Your reviews consistently rate lower than peers (z-score: {severity_bias_zscore:.1f}). While high standards are good, consider if you're being overly critical. This extreme pattern has resulted in a -1.5 point adjustment."
                else:
                    bias_feedback = f" <b>Rating Observation:</b> Your reviews tend to rate lower than peers (avg difference: {severity_bias:.1f}). High standards are valuable - just ensure you're recognizing good work when present. No penalty applied."
        
        # Base reliability message with more encouraging tone
        base_message = ""
        if pd.isna(score) or score == 'N/A':
            base_message = "Your peer review completion has been recorded. Insufficient data for detailed reliability analysis."
        elif score >= 85:
            base_message = "Outstanding work! Your reviews demonstrate excellent reliability and strong alignment with peer consensus. You're contributing meaningfully to the peer learning process."
        elif score >= 75:
            base_message = "Excellent work! Your reviews show strong reliability and good consistency with peer consensus. You demonstrate solid understanding of the evaluation criteria."
        elif score >= 65:
            base_message = "Good work! Your reviews show reasonable consistency with peer consensus. You're developing strong evaluation skills with good understanding of the criteria."
        elif score >= 55:
            base_message = "Satisfactory work! Your reviews show moderate consistency with peers. Continue focusing on the evaluation criteria to strengthen your assessment alignment."
        elif score >= 45:
            base_message = "Your reviews show some consistency challenges. Consider using more of the rating scale range and referring back to the evaluation guidelines to better distinguish performance levels."
        else:
            base_message = "Your reviews show significant variance from peer consensus. Please review the evaluation guidelines carefully and consider how to better align with the assessment criteria."
        
        # Support message for lower scores
        support_message = ""
        if score < 55 or (not pd.isna(severity_bias_zscore) and abs(severity_bias_zscore) > 2.0) or time_category in ["very_rushed", "technical_issue"]:
            support_message = " If you have questions about this assessment or need clarification on the evaluation criteria, please contact me."
        
        return base_message + time_feedback + bias_feedback + support_message

    student_feedback_sections = []
    if not reviewer_scores.empty:
        # Calculate total time per reviewer and count reviews
        reviewer_stats = merged.groupby('email_clean').agg({
            'DurationMin': 'sum',
            'Group#_reviewing': 'count'
        }).reset_index()
        reviewer_stats.columns = ['email_clean', 'total_time_minutes', 'num_reviews']
        
        # Calculate z-scores for severity bias to identify extreme outliers
        if not reviewer_scores.empty and 'severity_bias' in reviewer_scores.columns:
            severity_mean = reviewer_scores['severity_bias'].mean()
            severity_std = reviewer_scores['severity_bias'].std()
            if severity_std > 0:
                reviewer_scores['severity_bias_zscore'] = (reviewer_scores['severity_bias'] - severity_mean) / severity_std
            else:
                reviewer_scores['severity_bias_zscore'] = 0
        
        # Merge with roster to get student IDs and names
        feedback_data = reviewer_scores.merge(
            roster[['email_clean', 'student_id_clean']],
            left_on='reviewer_email',
            right_on='email_clean',
            how='left'
        )
        
        # Add reviewer statistics
        feedback_data = feedback_data.merge(
            reviewer_stats,
            left_on='reviewer_email',
            right_on='email_clean',
            how='left'
        )
        
        # Get student names if available in roster
        student_names = {}
        name_columns = [col for col in roster.columns if 'name' in col.lower()]
        if len(name_columns) >= 2:  # FirstName and LastName
            first_name_col = next((col for col in name_columns if 'first' in col.lower()), name_columns[0])
            last_name_col = next((col for col in name_columns if 'last' in col.lower()), name_columns[1])
            for _, row in roster.iterrows():
                full_name = f"{row[first_name_col]} {row[last_name_col]}"
                student_names[row['email_clean']] = full_name
        elif len(name_columns) == 1:  # Single Name column
            for _, row in roster.iterrows():
                student_names[row['email_clean']] = row[name_columns[0]]
        
        for _, row in feedback_data.iterrows():
            score = row['reliability_score']
            total_time = row.get('total_time_minutes', 0)
            num_reviews = row.get('num_reviews', 0)
            severity_bias = row.get('severity_bias', 0)
            severity_bias_zscore = row.get('severity_bias_zscore', 0)
            
            grade, time_category, time_adjustment = reliability_score_to_grade(score, total_time, num_reviews, severity_bias_zscore)
            grade = moderate_peer_review_completion_grade(grade, num_reviews)
            feedback = get_feedback_message(score, grade, total_time, num_reviews, severity_bias, severity_bias_zscore, time_category, time_adjustment)
            student_id = row.get('student_id_clean', row['reviewer_email'])
            student_name = student_names.get(row['reviewer_email'], 'Name not available')
            
            # Determine special circumstances
            special_notes = []
            if num_reviews < 2:
                special_notes.append(f"<strong>Incomplete:</strong> Only {num_reviews} review(s) completed (required: 2)")
            elif total_time < 2.0:
                special_notes.append(f"<strong>Time Warning:</strong> Total time for 2 reviews: {total_time:.1f} minutes")
            
            if not pd.isna(severity_bias_zscore) and abs(severity_bias_zscore) > 2.0:
                special_notes.append(f"<strong>Extreme Bias:</strong> Rating pattern {severity_bias_zscore:.1f} standard deviations from peer average")
            elif not pd.isna(severity_bias) and abs(severity_bias) > 0.5:
                bias_type = "higher" if severity_bias > 0 else "lower"
                special_notes.append(f"<strong>Note:</strong> Tends to rate {bias_type} than peers (not penalized)")
            
            special_notes_html = ""
            if special_notes:
                special_notes_html = "<p>" + "<br>".join(special_notes) + "</p>"
            
            feedback_section = f"""
    <!-- STUDENT FEEDBACK: {student_id} -->
    <div class="student-feedback" style="border: 2px dashed #007acc; margin: 20px 0; padding: 15px; background-color: #f8f9fa;">
        <h3>Student: {student_name} ({student_id})</h3>
        <p><strong>Peer Review Quality Grade: {grade}/10</strong></p>
        <p><strong>Reliability Score: {score:.1f}/100</strong></p>
        <p><strong>Reviews Completed: {num_reviews}/2</strong></p>
        <p><strong>Total Time for Both Reviews: {total_time:.1f} minutes</strong></p>
        {special_notes_html}
        <p><strong>Feedback:</strong> {feedback}</p>
        <hr style="border: 1px solid #007acc;">
        <p><em>Copy the section above (including grade and feedback) to paste into D2L for this student.</em></p>
    </div>
    """
            student_feedback_sections.append(feedback_section)
    
    html_student_feedback = '\n'.join(student_feedback_sections)

    # ── Reviewer quality summary table ──────────────────────────────
    html_reviewer_quality = ""
    if not reviewer_scores.empty:
        quality_summary = reviewer_scores[[
            'reviewer_email', 'reliability_score', 'consistency_score',
            'severity_bias', 'cronbach_alpha', 'icc_value'
        ]].copy()
        
        # Add student IDs
        quality_summary = quality_summary.merge(
            roster[['email_clean', 'student_id_clean']],
            left_on='reviewer_email',
            right_on='email_clean',
            how='left'
        )
        
        quality_display = quality_summary[[
            'student_id_clean', 'reliability_score', 'consistency_score',
            'severity_bias', 'cronbach_alpha', 'icc_value'
        ]].round(3)
        
        quality_display = quality_display.rename(columns={
            'student_id_clean': 'Student ID',
            'reliability_score': 'Overall Score',
            'consistency_score': 'Consistency',
            'severity_bias': 'Severity Bias',
            'cronbach_alpha': 'Cronbach α',
            'icc_value': 'ICC'
        })
        
        html_reviewer_quality = quality_display.to_html(
            index=False, classes="table", border=0,
            justify="left", float_format="%.3f"
        )

    # ── Generate instructor follow-up section for short completion times ───
    def generate_instructor_followup_section(reviewer_scores, time_data):
        """Generate a section highlighting students who need personal follow-up."""
        
        # Identify students needing follow-up based on extremely short times
        followup_needed = []
        
        for _, row in time_data.iterrows():
            total_time = row.get('DurationMin', 0)
            if pd.isna(total_time):
                continue
                
            # Group by email to get total time for all reviews
            student_time = time_data[time_data['email_clean'] == row['email_clean']]['DurationMin'].sum()
            
            if student_time < 0.5:  # Less than 30 seconds total
                student_reviews = time_data[time_data['email_clean'] == row['email_clean']]
                num_reviews = len(student_reviews)
                
                # Get reliability score if available
                reliability_score = 'N/A'
                grade = 'N/A'
                if not reviewer_scores.empty:
                    student_data = reviewer_scores[reviewer_scores['reviewer_email'] == row['email_clean']]
                    if not student_data.empty:
                        reliability_score = student_data.iloc[0].get('reliability_score', 'N/A')
                        grade = student_data.iloc[0].get('grade', 'N/A')
                
                followup_needed.append({
                    'email': row['email_clean'],
                    'total_time': student_time,
                    'num_reviews': num_reviews,
                    'reliability_score': reliability_score,
                    'grade': grade,
                    'reason': 'Extremely short completion time - possible technical issue'
                })
        
        # Remove duplicates (since we might have multiple rows per student)
        followup_needed = {item['email']: item for item in followup_needed}.values()
        
        if not followup_needed:
            return "<h2>Student Follow-up Status</h2><p><strong>✓ No students requiring immediate follow-up</strong> - All students had reasonable completion times.</p>"
        
        html = "<h2>Students Requiring Personal Follow-up</h2>"
        html += "<div style='background-color: #fff3cd; padding: 15px; margin: 10px 0; border-left: 4px solid #ffc107;'>"
        html += "<p><strong>⚠️ The following students had unusually short completion times and should be contacted personally:</strong></p>"
        html += "<table class='table'>"
        html += "<tr><th>Student Email</th><th>Total Time (min)</th><th>Reviews Completed</th><th>Current Grade</th><th>Suggested Action</th></tr>"
        
        for student in followup_needed:
            html += f"<tr style='background-color: #ffeaa7;'>"
            html += f"<td>{student['email']}</td>"
            html += f"<td>{student['total_time']:.1f} ({student['total_time']*60:.0f} sec)</td>"
            html += f"<td>{student['num_reviews']}</td>"
            html += f"<td>{student['grade']}/10</td>"
            html += f"<td>Contact for possible technical issues</td>"
            html += f"</tr>"
        
        html += "</table>"
        html += "<div style='background-color: #e7f3ff; padding: 10px; margin-top: 15px; border: 1px solid #007acc; border-radius: 5px;'>"
        html += "<p><strong>📧 Suggested Email Template:</strong></p>"
        html += "<div style='background-color: white; padding: 10px; border-left: 3px solid #007acc; font-style: italic;'>"
        html += "<p>Hi [Student Name],</p>"
        html += "<p>I noticed your peer review completion time was quite short (under 30 seconds). This might indicate a technical issue with the form or browser.</p>"
        html += "<p>If you experienced any difficulties, please let me know what they were and I'll adjust your grade. - I just want to ensure everyone gets fair assessment.</p>"
        html += "<p>Best regards,<br>[Your name]</p>"
        html += "</div></div></div>"
        
        return html

    # Generate the follow-up section
    followup_html = generate_instructor_followup_section(reviewer_scores, time_data)

    # ── Enhanced Jinja2 template ────────────────────────────────────
    tmpl = Template(
        """
        <html><head><title>Instructor Report</title>
        <style>
          body{font-family:sans-serif;max-width:1000px;margin:0 auto;padding:20px}
          .table{border-collapse:collapse;width:100%;margin-bottom:20px}
          .table th,.table td{border:1px solid #ddd;padding:8px;text-align:left}
          .table th{background:#f2f2f2;font-weight:bold}
          .metric-description{background:#f9f9f9;padding:15px;margin:10px 0;border-left:4px solid #007acc}
          .section{margin-bottom:30px}
          .reliability-high{background-color:#d4edda}
          .reliability-medium{background-color:#fff3cd}
          .reliability-low{background-color:#f8d7da}
          .student-feedback{border:2px dashed #007acc;margin:20px 0;padding:15px;background-color:#f8f9fa}
          h1{color:#333;border-bottom:2px solid #007acc;padding-bottom:10px}
          h2{color:#555;margin-top:30px}
          .feedback-instructions{background:#e7f3ff;padding:15px;margin:20px 0;border:1px solid #007acc;border-radius:5px}
        </style></head><body>
        
        <h1>Instructor Report</h1>
        
        <div class="metric-description">
            <strong>Reliability Metrics Explanation:</strong><br>
            • <strong>Cronbach's α</strong>: Internal consistency (>0.7 good, >0.8 excellent)<br>
            • <strong>ICC</strong>: Inter-rater reliability (>0.75 good, >0.9 excellent)<br>
            • <strong>Reliability Score</strong>: Composite score (0-100, higher = more reliable reviewer)<br>
            • <strong>Consistency</strong>: Agreement with group consensus (higher = better)<br>
            • <strong>Severity Bias</strong>: Tendency to rate higher/lower than peers (closer to 0 = better)
        </div>

        <div class="section">
            <h2>Group Results with Reliability Metrics</h2>
            {{ group_html | safe }}
        </div>

        {{ followup_html | safe }}

        {% if reviewer_quality_html %}
        <div class="section">
            <h2>Reviewer Quality Assessment</h2>
            <p><em>Ranked by overall reliability score. Higher scores indicate more reliable reviewers.</em></p>
            {{ reviewer_quality_html | safe }}
        </div>
        {% endif %}

        {% if one_only %}
        <div class="section">
            <h2>Exactly One Submission ({{ one_only|length }})</h2>
            <p>{{ one_only|join(", ") }}</p>
        </div>
        {% endif %}

        {% if three_plus %}
        <div class="section">
            <h2>Three+ Submissions ({{ three_plus|length }})</h2>
            <p>{{ three_plus|join(", ") }}</p>
        </div>
        {% endif %}

        {% if no_submit %}
        <div class="section">
            <h2>No Submission Received ({{ no_submit|length }})</h2>
            <p>{{ no_submit|join(", ") }}</p>
        </div>
        {% endif %}

        <div class="section">
            <h2>Time on Task and Reliability Scores</h2>
            <p><em>Each review duration in minutes with corresponding reliability assessment.</em></p>
            {{ time_html | safe }}
        </div>

        {% if student_feedback_html %}
        <div class="section">
            <h2>Individual Student Feedback (Copy to D2L)</h2>
            <div class="feedback-instructions">
                <strong>Instructions:</strong> Each dashed box below contains feedback for one student. 
                Copy the content (including grade and feedback text) and paste directly into the student's 
                feedback section in D2L. The grades are out of 10 based on peer review quality.
            </div>
            {{ student_feedback_html | safe }}
        </div>
        {% endif %}

        <div class="metric-description">
            <strong>Interpretation Guidelines:</strong><br>
            • Groups with Cronbach's α < 0.6 may have inconsistent rating criteria<br>
            • ICC < 0.5 suggests poor inter-rater agreement<br>
            • High severity bias may indicate need for calibration training<br>
            • Very low variance (halo score) may suggest insufficient differentiation<br><br>
            <strong>Grading Scale (out of 10):</strong><br>
            • 10: Reliability Score 85+ (Excellent)<br>
            • 9: Reliability Score 75-84 (Very Good)<br>
            • 8: Reliability Score 65-74 (Good)<br>
            • 7: Reliability Score 55-64 (Satisfactory)<br>
            • 6: Reliability Score 45-54 (Needs Improvement)<br>
            • 5: Reliability Score 35-44 (Poor)<br>
            • 4: Reliability Score <35 (Very Poor)<br><br>
            <strong>Grading Adjustments (Graduated Time System):</strong><br>
            • <span style="color: green;">Exceptional thoroughness (>20 minutes total): +1.0 bonus point</span><br>
            • <span style="color: green;">Good thoroughness (15-20 minutes total): +0.5 bonus point</span><br>
            • Adequate time (10-15 minutes total): No adjustment<br>
            • Somewhat rushed (5-10 minutes total): -0.5 points<br>
            • Rushed (2-5 minutes total): -1.0 point<br>
            • Very rushed (<2 minutes total): -2.0 points<br>
            • Extremely short (<30 seconds total): -3.0 points + instructor follow-up<br>
            • Incomplete reviews (<2 required): Half marks<br>
            • Extreme bias (>2 standard deviations from peer average): -1.5 points<br>
            • Minor bias variations are noted but not penalized<br><br>
            <strong>Time Measurement:</strong> Total time refers to combined time for both required peer reviews.
        </div>

        </body></html>
        """
    )

    return tmpl.render(
        group_html=html_group,
        time_html=html_time,
        reviewer_quality_html=html_reviewer_quality,
        student_feedback_html=html_student_feedback,
        followup_html=followup_html,
        one_only=tag(one_only),
        three_plus=tag(three_plus),
        no_submit=tag(no_submit),
    )


# ──────────────────────────────── Main ──────────────────────────────── #

def cli() -> None:
    ap = argparse.ArgumentParser(
        description="Generate enhanced peer‑evaluation instructor report with reliability metrics."
    )
    ap.add_argument("reviews", type=Path, help="Excel file with peer reviews")
    ap.add_argument("roster", type=Path, help="Excel class‑list file")
    ap.add_argument("-o", "--out", type=Path, default="Instructor_Report.html",
                    help="HTML output path")
    args = ap.parse_args()

    LOG.info("Loading data …")
    review_df = read_review(args.reviews)
    roster_df = read_roster(args.roster)

    LOG.info("Building enhanced report with reliability metrics …")
    html = build_enhanced_report(review_df, roster_df)

    args.out.write_text(html, encoding="utf‑8")
    LOG.info("Instructor report saved to %s", args.out)


if __name__ == "__main__":
    cli()
